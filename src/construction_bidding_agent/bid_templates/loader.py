"""静态投标模板加载。"""

from __future__ import annotations

import json
import re
from pathlib import Path
from typing import Any


def load_bid_templates(template_dir: str | Path) -> list[dict[str, Any]]:
    root = Path(template_dir)
    if not root.exists():
        return []
    templates = []
    for path in sorted(root.glob("*.json")):
        try:
            data = json.loads(path.read_text(encoding="utf-8"))
        except Exception:
            continue
        data.setdefault("template_id", path.stem)
        data.setdefault("source_file", path.name)
        data.setdefault("chapter_count", len(data.get("chapters") or []))
        data.setdefault("table_count", len(data.get("tables") or data.get("table_templates") or []))
        data.setdefault("usage_boundary", "模板只做推荐和预览，不自动覆盖已确认目录或正文。")
        templates.append(data)
    return templates


def save_bid_template(template_dir: str | Path, template: dict[str, Any]) -> dict[str, Any]:
    root = Path(template_dir)
    root.mkdir(parents=True, exist_ok=True)
    normalized = normalize_bid_template(template)
    target = root / f"{normalized['template_id']}.json"
    target.write_text(json.dumps(normalized, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")
    normalized["source_file"] = target.name
    return normalized


def normalize_bid_template(template: dict[str, Any]) -> dict[str, Any]:
    data = dict(template or {})
    name = str(data.get("name") or data.get("title") or "").strip()
    if not name:
        raise ValueError("模板名称不能为空。")
    template_id = str(data.get("template_id") or "").strip() or _slugify(name)
    chapters = [_normalize_chapter(item) for item in _list(data.get("chapters"))]
    tables = [str(item).strip() for item in _list(data.get("tables") or data.get("table_templates")) if str(item).strip()]
    scenarios = [str(item).strip() for item in _list(data.get("applicable_scenarios")) if str(item).strip()]
    tags = [str(item).strip() for item in _list(data.get("tags")) if str(item).strip()]
    normalized = {
        **data,
        "template_id": _slugify(template_id),
        "name": name,
        "version": str(data.get("version") or "v1").strip() or "v1",
        "project_type": str(data.get("project_type") or "construction").strip() or "construction",
        "description": str(data.get("description") or "企业上传的投标模板。").strip(),
        "tags": tags,
        "applicable_scenarios": scenarios,
        "usage_boundary": str(data.get("usage_boundary") or "模板只做推荐和预览，不自动覆盖已确认目录或正文。").strip(),
        "chapters": chapters,
        "tables": tables,
        "chapter_count": len(chapters),
        "table_count": len(tables),
    }
    return normalized


def parse_bid_template_docx(
    docx_path: str | Path,
    *,
    name: str | None = None,
    project_type: str = "construction",
    version: str = "v1",
    description: str | None = None,
) -> dict[str, Any]:
    from docx import Document

    path = Path(docx_path)
    doc = Document(str(path))
    paragraphs = [paragraph.text.strip() for paragraph in doc.paragraphs if paragraph.text and paragraph.text.strip()]
    title = name or _docx_title(paragraphs, path)
    chapters = _chapters_from_docx_paragraphs(paragraphs)
    tables = _tables_from_docx(doc)
    tags = _infer_tags(" ".join([title, *paragraphs[:80], *tables]))
    return normalize_bid_template(
        {
            "template_id": _slugify(title),
            "name": title,
            "version": version or "v1",
            "project_type": project_type or "construction",
            "description": description or f"由 {path.name} 自动解析生成的投标模板。",
            "tags": tags,
            "applicable_scenarios": [f"{_project_type_label(project_type)}项目技术标编制参考"],
            "usage_boundary": "由 Word 模板解析生成，仅作为目录、章节重点和表格清单参考，不自动覆盖项目成果。",
            "chapters": chapters,
            "tables": tables,
            "parse_summary": {
                "source_file": path.name,
                "paragraph_count": len(paragraphs),
                "detected_chapter_count": len(chapters),
                "detected_table_count": len(tables),
            },
        }
    )


def _chapters_from_docx_paragraphs(paragraphs: list[str]) -> list[dict[str, Any]]:
    heading_indexes: list[tuple[int, str]] = []
    for index, text in enumerate(paragraphs):
        clean = _clean_heading(text)
        if _looks_like_heading(clean):
            heading_indexes.append((index, clean))
    if not heading_indexes:
        heading_indexes = [(index, _clean_heading(text)) for index, text in enumerate(paragraphs[:8]) if len(text) <= 32]
    chapters: list[dict[str, Any]] = []
    for pos, (index, title) in enumerate(heading_indexes[:20]):
        next_index = heading_indexes[pos + 1][0] if pos + 1 < len(heading_indexes) else min(len(paragraphs), index + 8)
        focus = _writing_focus_from_text(paragraphs[index + 1 : next_index])
        chapters.append({"title": title, "writing_focus": focus})
    return chapters


def _tables_from_docx(doc: Any) -> list[str]:
    tables: list[str] = []
    for index, table in enumerate(doc.tables, start=1):
        first_row = []
        try:
            first_row = [cell.text.strip() for cell in table.rows[0].cells if cell.text.strip()]
        except Exception:
            first_row = []
        label = " / ".join(first_row[:4]) if first_row else f"表格 {index}"
        tables.append(label[:80])
    return tables[:20]


def _writing_focus_from_text(lines: list[str]) -> list[str]:
    focus: list[str] = []
    for line in lines:
        clean = re.sub(r"^\s*[\d一二三四五六七八九十]+[、.．\s-]*", "", line).strip()
        if 4 <= len(clean) <= 42 and clean not in focus:
            focus.append(clean)
        if len(focus) >= 4:
            break
    return focus or ["结合招标评分点补充项目化内容", "保持企业标准表达", "复核与当前项目参数一致性"]


def _docx_title(paragraphs: list[str], path: Path) -> str:
    for text in paragraphs[:8]:
        clean = _clean_heading(text)
        if 4 <= len(clean) <= 40:
            return clean
    return path.stem


def _clean_heading(text: str) -> str:
    return re.sub(r"\s+", " ", str(text or "")).strip(" ：:;；")


def _looks_like_heading(text: str) -> bool:
    if not text or len(text) > 42:
        return False
    patterns = [
        r"^第[一二三四五六七八九十\d]+[章节篇部分]",
        r"^[一二三四五六七八九十]+[、.．]",
        r"^\d+(\.\d+){0,3}[、.．\s]",
    ]
    if any(re.search(pattern, text) for pattern in patterns):
        return True
    return any(token in text for token in ["施工", "质量", "安全", "进度", "文明", "管理", "措施", "方案", "部署", "组织"]) and len(text) <= 24


def _infer_tags(text: str) -> list[str]:
    seeds = ["施工", "质量", "安全", "进度", "文明施工", "环保", "BIM", "EPC", "设计", "采购", "表格"]
    tags = [seed for seed in seeds if seed.lower() in text.lower()]
    return tags[:6] or ["企业模板"]


def _slugify(value: str) -> str:
    text = re.sub(r"[^0-9a-zA-Z\u4e00-\u9fa5_-]+", "_", str(value or "").strip()).strip("_").lower()
    if not text:
        text = "bid_template"
    return text[:80]


def _project_type_label(value: str) -> str:
    return {"construction": "施工总承包", "epc": "EPC", "general": "通用"}.get(value or "", "建设工程")


def _normalize_chapter(item: Any) -> dict[str, Any]:
    if isinstance(item, dict):
        title = str(item.get("title") or item.get("name") or "未命名章节").strip()
        focus = [str(value).strip() for value in _list(item.get("writing_focus")) if str(value).strip()]
        return {"title": title, "writing_focus": focus}
    return {"title": str(item or "未命名章节").strip(), "writing_focus": []}


def _list(value: Any) -> list[Any]:
    return value if isinstance(value, list) else []
