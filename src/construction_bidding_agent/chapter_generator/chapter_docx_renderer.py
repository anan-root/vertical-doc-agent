"""技术标章节初稿 Word 渲染器。

章节生成阶段输出结构化 JSON，本模块只负责把正文、表格、图片引用和项目专属图片占位
渲染为可编辑的 DOCX 样稿。它不再调用 LLM，也不改写正文内容。
"""

from __future__ import annotations

import json
import re
import struct
import time
import zipfile
from dataclasses import dataclass, field
from io import BytesIO
from pathlib import Path
from typing import Any

from docx import Document
from docx.document import Document as DocumentObject
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.image.exceptions import UnrecognizedImageError
from docx.opc.constants import CONTENT_TYPE as CT
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.opc.part import Part
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.oxml.shape import CT_Inline
from docx.shared import Cm, Pt, RGBColor, Twips
from PIL import Image, UnidentifiedImageError

from .word_export_profile import load_word_export_profile, merge_word_export_profile


DEFAULT_LIBRARY = Path("outputs/json/excellent_bid_material_library_two_word_sources.json")
DEFAULT_RAW_ROOT = Path("data/raw")
DEFAULT_RENDER_PROFILE = Path("configs/docx-render-profile.json")
PAGE_WIDTH_CM = 21.0
PAGE_HEIGHT_CM = 29.7
MARGIN_CM = 2.5
BODY_WIDTH_CM = PAGE_WIDTH_CM - MARGIN_CM * 2
IMAGE_MAX_WIDTH_CM = 7.6
FULL_IMAGE_MAX_WIDTH_CM = 12.8
FULL_IMAGE_MAX_HEIGHT_CM = 8.5
GRID_IMAGE_MAX_HEIGHT_CM = 5.6
GRID_IMAGE_CELL_PADDING_CM = 0.45
FINAL_DOCX_MODE = "final"
REVIEW_DOCX_MODE = "review"
TITLE_FONT = "宋体"
BODY_FONT = "宋体"
BODY_FONT_SIZE_PT = 12
TABLE_FONT_SIZE_PT = 12
TITLE_INDENT_CM = 0
BODY_FIRST_LINE_INDENT_CM = 0.74

HIGH_DETAIL_IMAGE_KEYWORDS = [
    "节点",
    "详图",
    "大样",
    "剖面",
    "平立面",
    "平面图",
    "立面图",
    "尺寸",
    "标注",
    "布置图",
    "网络图",
    "设计图",
    "配筋图",
]
MEDIUM_DETAIL_IMAGE_KEYWORDS = [
    "示意图",
    "构造",
    "流程",
    "控制网",
    "轴线",
    "模板",
    "钢筋",
    "架体",
    "支设",
    "搭设",
    "做法图",
]
PHOTO_IMAGE_KEYWORDS = [
    "照片",
    "现场",
    "实景",
    "样板",
    "成品",
    "浇筑",
    "振捣",
    "抹面",
    "养护",
    "砌筑",
    "防水",
    "涂刷",
]
VECTOR_IMAGE_CONTENT_TYPES = {
    ".emf": CT.X_EMF,
    ".wmf": CT.X_WMF,
}


@dataclass(slots=True)
class ImageLayoutProfile:
    body_width_cm: float = BODY_WIDTH_CM
    single_image_max_width_cm: float = FULL_IMAGE_MAX_WIDTH_CM
    single_image_max_height_cm: float = FULL_IMAGE_MAX_HEIGHT_CM
    grid_cell_max_width_cm: float = IMAGE_MAX_WIDTH_CM
    grid_cell_padding_cm: float = GRID_IMAGE_CELL_PADDING_CM
    one_column_max_height_cm: float = FULL_IMAGE_MAX_HEIGHT_CM
    two_column_max_height_cm: float = 6.4
    three_column_max_height_cm: float = GRID_IMAGE_MAX_HEIGHT_CM
    high_detail_keywords: list[str] = field(default_factory=lambda: HIGH_DETAIL_IMAGE_KEYWORDS.copy())
    medium_detail_keywords: list[str] = field(default_factory=lambda: MEDIUM_DETAIL_IMAGE_KEYWORDS.copy())
    photo_keywords: list[str] = field(default_factory=lambda: PHOTO_IMAGE_KEYWORDS.copy())


@dataclass(slots=True)
class ImageSource:
    source_id: str
    paths: list[Path] = field(default_factory=list)


@dataclass(slots=True)
class ImageLayoutItem:
    block: dict[str, Any]
    image_bytes: bytes | None
    max_columns: int


@dataclass(slots=True)
class RenderStats:
    chapter_count: int = 0
    heading_count: int = 0
    heading1_count: int = 0
    heading2_count: int = 0
    heading3_count: int = 0
    paragraph_count: int = 0
    table_count: int = 0
    image_ref_count: int = 0
    rendered_image_count: int = 0
    missing_image_count: int = 0
    placeholder_count: int = 0
    image_layout_one_column_row_count: int = 0
    image_layout_two_column_row_count: int = 0
    image_layout_three_column_row_count: int = 0
    image_processing_duration_seconds: float = 0.0

    def to_dict(self) -> dict[str, Any]:
        return {
            "chapter_count": self.chapter_count,
            "heading_count": self.heading_count,
            "heading1_count": self.heading1_count,
            "heading2_count": self.heading2_count,
            "heading3_count": self.heading3_count,
            "paragraph_count": self.paragraph_count,
            "table_count": self.table_count,
            "image_ref_count": self.image_ref_count,
            "rendered_image_count": self.rendered_image_count,
            "missing_image_count": self.missing_image_count,
            "placeholder_count": self.placeholder_count,
            "image_layout_one_column_row_count": self.image_layout_one_column_row_count,
            "image_layout_two_column_row_count": self.image_layout_two_column_row_count,
            "image_layout_three_column_row_count": self.image_layout_three_column_row_count,
            "image_processing_duration_seconds": round(self.image_processing_duration_seconds, 4),
        }


def render_chapter_docx_from_file(
    generation_result_json: str | Path,
    output_docx: str | Path,
    *,
    material_library_json: str | Path | None = DEFAULT_LIBRARY,
    raw_root: str | Path = DEFAULT_RAW_ROOT,
    render_profile_json: str | Path | None = DEFAULT_RENDER_PROFILE,
    word_export_profile: dict[str, Any] | str | Path | None = None,
    title: str | None = None,
    output_mode: str = REVIEW_DOCX_MODE,
) -> dict[str, int]:
    """从章节生成结果 JSON 渲染 Word 初稿，并返回渲染统计。"""

    data = json.loads(Path(generation_result_json).read_text(encoding="utf-8"))
    return write_chapter_docx(
        data,
        output_docx,
        material_library_json=material_library_json,
        raw_root=raw_root,
        render_profile_json=render_profile_json,
        word_export_profile=word_export_profile,
        title=title,
        output_mode=output_mode,
    )


def write_chapter_docx(
    generation_result: dict[str, Any],
    output_docx: str | Path,
    *,
    material_library_json: str | Path | None = DEFAULT_LIBRARY,
    raw_root: str | Path = DEFAULT_RAW_ROOT,
    render_profile_json: str | Path | None = DEFAULT_RENDER_PROFILE,
    word_export_profile: dict[str, Any] | str | Path | None = None,
    title: str | None = None,
    output_mode: str = REVIEW_DOCX_MODE,
) -> dict[str, int]:
    """写入章节 Word 初稿。"""

    target = Path(output_docx)
    target.parent.mkdir(parents=True, exist_ok=True)
    export_profile = _load_word_export_profile(word_export_profile)
    doc = Document()
    _setup_document(doc, export_profile)
    profile = _load_image_layout_profile(render_profile_json, export_profile)
    resolver = _build_image_resolver(material_library_json, raw_root)
    chapters = [chapter for chapter in generation_result.get("chapters") or [] if isinstance(chapter, dict)]
    stats = RenderStats(chapter_count=len(chapters))

    mode = _normalize_output_mode(output_mode)
    if mode == REVIEW_DOCX_MODE:
        _render_front_matter(doc, title or "技术标章节初稿", generation_result, export_profile)
    else:
        _render_toc_page(doc, export_profile)
    for index, chapter in enumerate(chapters, start=1):
        if mode == REVIEW_DOCX_MODE:
            doc.add_section(WD_SECTION_START.NEW_PAGE)
        _render_chapter(
            doc,
            chapter,
            index,
            resolver,
            stats,
            profile,
            export_profile,
            output_mode=mode,
        )

    doc.save(target)
    return stats.to_dict()


def _load_word_export_profile(profile: dict[str, Any] | str | Path | None) -> dict[str, Any]:
    if isinstance(profile, (str, Path)):
        return load_word_export_profile(profile)
    return merge_word_export_profile(profile if isinstance(profile, dict) else None)


def _setup_document(doc: DocumentObject, export_profile: dict[str, Any]) -> None:
    section = doc.sections[0]
    _apply_section_page_setup(section, export_profile)
    styles = doc.styles
    normal = styles["Normal"]
    body_profile = export_profile["body"]
    _apply_style_font(normal, body_profile)
    _apply_style_paragraph(normal, body_profile)
    heading_specs = {
        "Title": {
            "font_family": export_profile["heading_1"].get("font_family", TITLE_FONT),
            "font_size_pt": 20,
            "bold": True,
            "color": "000000",
            "alignment": "center",
            "first_line_indent_chars": 0,
            "line_spacing": export_profile["body"].get("line_spacing", 1.35),
            "space_before_pt": 0,
            "space_after_pt": 12,
        },
        "Heading 1": export_profile["heading_1"],
        "Heading 2": export_profile["heading_2"],
        "Heading 3": export_profile["heading_3"],
    }
    for style_name, style_profile in heading_specs.items():
        style = styles[style_name]
        _apply_style_font(style, style_profile)
        _apply_style_paragraph(style, style_profile)
    styles["Heading 1"].paragraph_format.page_break_before = bool(export_profile["heading_1"].get("page_break_before"))


def _apply_section_page_setup(section: Any, export_profile: dict[str, Any]) -> None:
    page = export_profile["page"]
    orientation = str(page.get("orientation") or "portrait")
    if orientation == "landscape":
        section.page_width = Cm(PAGE_HEIGHT_CM)
        section.page_height = Cm(PAGE_WIDTH_CM)
    else:
        section.page_width = Cm(PAGE_WIDTH_CM)
        section.page_height = Cm(PAGE_HEIGHT_CM)
    section.top_margin = Cm(float(page.get("top_margin_cm") or MARGIN_CM))
    section.bottom_margin = Cm(float(page.get("bottom_margin_cm") or MARGIN_CM))
    section.left_margin = Cm(float(page.get("left_margin_cm") or MARGIN_CM))
    section.right_margin = Cm(float(page.get("right_margin_cm") or MARGIN_CM))
    section.header_distance = Cm(float(page.get("header_distance_cm") or 1.5))
    section.footer_distance = Cm(float(page.get("footer_distance_cm") or 1.75))


def _body_width_cm(export_profile: dict[str, Any] | None = None) -> float:
    if not export_profile:
        return BODY_WIDTH_CM
    page = export_profile["page"]
    orientation = str(page.get("orientation") or "portrait")
    page_width = PAGE_HEIGHT_CM if orientation == "landscape" else PAGE_WIDTH_CM
    return max(
        4.0,
        page_width - float(page.get("left_margin_cm") or MARGIN_CM) - float(page.get("right_margin_cm") or MARGIN_CM),
    )


def _apply_style_font(style: Any, style_profile: dict[str, Any]) -> None:
    font_family = str(style_profile.get("font_family") or BODY_FONT)
    style.font.name = font_family
    style._element.rPr.rFonts.set(qn("w:eastAsia"), font_family)
    style.font.size = Pt(float(style_profile.get("font_size_pt") or BODY_FONT_SIZE_PT))
    style.font.bold = bool(style_profile.get("bold", False))
    style.font.color.rgb = _rgb_color(str(style_profile.get("color") or "000000"))


def _apply_style_paragraph(style: Any, style_profile: dict[str, Any]) -> None:
    _apply_paragraph_format(style.paragraph_format, style_profile)
    style.paragraph_format.alignment = _paragraph_alignment(str(style_profile.get("alignment") or "left"))


def _apply_paragraph_format(paragraph_format: Any, style_profile: dict[str, Any]) -> None:
    paragraph_format.line_spacing = float(style_profile.get("line_spacing") or 1.35)
    paragraph_format.first_line_indent = Cm(_indent_chars_to_cm(style_profile.get("first_line_indent_chars") or 0))
    paragraph_format.space_before = Pt(float(style_profile.get("space_before_pt") or 0))
    paragraph_format.space_after = Pt(float(style_profile.get("space_after_pt") or 0))


def _apply_run_font(run: Any, style_profile: dict[str, Any], bold: bool | None = None) -> None:
    font_family = str(style_profile.get("font_family") or BODY_FONT)
    run.font.name = font_family
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_family)
    run.font.size = Pt(float(style_profile.get("font_size_pt") or BODY_FONT_SIZE_PT))
    run.bold = bool(style_profile.get("bold", False)) if bold is None else bool(bold)
    run.font.color.rgb = _rgb_color(str(style_profile.get("color") or "000000"))


def _caption_profile(export_profile: dict[str, Any]) -> dict[str, Any]:
    image_profile = export_profile["image"]
    return {
        "font_family": image_profile.get("caption_font_family") or BODY_FONT,
        "font_size_pt": image_profile.get("caption_font_size_pt") or 10.5,
        "bold": False,
        "color": "000000",
    }


def _paragraph_alignment(value: str) -> Any:
    return {
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
        "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
        "left": WD_ALIGN_PARAGRAPH.LEFT,
    }.get(str(value or "left").lower(), WD_ALIGN_PARAGRAPH.LEFT)


def _rgb_color(value: str) -> RGBColor:
    text = re.sub(r"[^0-9A-Fa-f]", "", str(value or "000000"))[:6].ljust(6, "0")
    return RGBColor(int(text[0:2], 16), int(text[2:4], 16), int(text[4:6], 16))


def _indent_chars_to_cm(value: Any) -> float:
    try:
        chars = float(value)
    except (TypeError, ValueError):
        chars = 0.0
    return max(0.0, chars * 0.37)


def _render_front_matter(
    doc: DocumentObject,
    title: str,
    generation_result: dict[str, Any],
    export_profile: dict[str, Any],
) -> None:
    heading = doc.add_paragraph(style="Title")
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    heading.add_run(title)
    meta = [
        ("生成模型", generation_result.get("model") or "-"),
        ("生成时间", generation_result.get("generated_at") or "-"),
        ("章节数量", str(len(generation_result.get("chapters") or []))),
    ]
    table = doc.add_table(rows=len(meta), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    body_width_cm = _body_width_cm(export_profile)
    _set_table_width(table, [4.0, body_width_cm - 4.0])
    for row, (key, value) in zip(table.rows, meta):
        row.cells[0].text = key
        row.cells[1].text = str(value)
        _shade_cell(row.cells[0], "F2F2F2")
        for cell in row.cells:
            _format_cell(cell, export_profile)
    doc.add_paragraph()


def _render_toc_page(doc: DocumentObject, export_profile: dict[str, Any]) -> None:
    toc_profile = export_profile["toc"]
    if not toc_profile.get("enabled", True):
        return
    title = str(toc_profile.get("title") or "目录")
    title_paragraph = doc.add_paragraph(style="Title")
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_paragraph.add_run(title)
    toc_paragraph = doc.add_paragraph()
    _add_toc_field(toc_paragraph, int(toc_profile.get("levels") or 3))
    if toc_profile.get("body_starts_new_page", True) or toc_profile.get("separate_page", True):
        body_section = doc.add_section(WD_SECTION_START.NEW_PAGE)
        _apply_section_page_setup(body_section, export_profile)
        if toc_profile.get("body_page_number_restart", True):
            _set_section_page_number_start(body_section, int(toc_profile.get("body_page_number_start") or 1))


def _add_toc_field(paragraph: Any, levels: int) -> None:
    levels = min(max(int(levels or 3), 1), 3)
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    run._r.append(begin)

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = f'TOC \\o "1-{levels}" \\h \\z \\u'
    run._r.append(instr)

    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    run._r.append(separate)
    fallback = OxmlElement("w:t")
    fallback.text = "请在 OnlyOffice 或 WPS 中更新目录"
    run._r.append(fallback)

    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(end)


def _set_section_page_number_start(section: Any, start: int) -> None:
    sect_pr = section._sectPr
    pg_num_type = sect_pr.find(qn("w:pgNumType"))
    if pg_num_type is None:
        pg_num_type = OxmlElement("w:pgNumType")
        sect_pr.append(pg_num_type)
    pg_num_type.set(qn("w:start"), str(max(1, int(start or 1))))


def _render_chapter(
    doc: DocumentObject,
    chapter: dict[str, Any],
    index: int,
    resolver: dict[tuple[str, str], bytes],
    stats: RenderStats,
    profile: ImageLayoutProfile,
    export_profile: dict[str, Any],
    *,
    output_mode: str,
) -> None:
    chapter_path = [str(part) for part in chapter.get("chapter_path") or [] if str(part).strip()]
    title = " > ".join(chapter_path) or str(chapter.get("title") or f"章节{index}")
    heading_number = str(index)
    _add_numbered_heading(doc, f"{heading_number}.{title}", level=1)
    _record_heading(stats, 1)
    check = chapter.get("score_response_check") or {}
    if output_mode == REVIEW_DOCX_MODE and check.get("response_summary"):
        paragraph = _add_body_paragraph(doc, export_profile)
        paragraph.add_run("评分点响应摘要：").bold = True
        paragraph.add_run(str(check.get("response_summary")))
    level2_counter = 0
    level3_counter = 0
    for section in chapter.get("sections") or []:
        if isinstance(section, dict):
            level = int(section.get("level") or 2)
            if level <= 2:
                level2_counter += 1
                level3_counter = 0
            else:
                level3_counter += 1
            if level <= 2:
                number = f"{heading_number}.{level2_counter}"
            else:
                number = f"{heading_number}.{level2_counter or 1}.{level3_counter}"
            _render_section(doc, section, resolver, stats, profile, export_profile, heading_number=number)
    if output_mode == REVIEW_DOCX_MODE:
        _render_review_items(doc, chapter.get("review_items") or [], export_profile)


def _render_section(
    doc: DocumentObject,
    section: dict[str, Any],
    resolver: dict[tuple[str, str], bytes],
    stats: RenderStats,
    profile: ImageLayoutProfile,
    export_profile: dict[str, Any],
    *,
    heading_number: str,
) -> None:
    heading = str(section.get("heading") or "未命名小节")
    level = int(section.get("level") or 2)
    rendered_level = 2 if level <= 2 else 3
    _add_numbered_heading(doc, f"{heading_number}.{heading}", level=rendered_level)
    _record_heading(stats, rendered_level)
    blocks = [block for block in section.get("blocks") or [] if isinstance(block, dict)]
    index = 0
    while index < len(blocks):
        block = blocks[index]
        block_type = str(block.get("type") or "")
        if block_type == "image_ref":
            group, index = _take_image_group(blocks, index)
            _render_image_group(doc, group, resolver, stats, profile, export_profile)
            continue
        _render_block(doc, block, resolver, stats, profile, export_profile)
        index += 1


def _render_block(
    doc: DocumentObject,
    block: dict[str, Any],
    resolver: dict[tuple[str, str], bytes],
    stats: RenderStats,
    profile: ImageLayoutProfile,
    export_profile: dict[str, Any],
) -> None:
    block_type = str(block.get("type") or "")
    if block_type == "paragraph":
        text = str(block.get("text") or "").strip()
        if text:
            paragraph = _add_body_paragraph(doc, export_profile)
            paragraph.add_run(text)
            _format_paragraph_font(paragraph, export_profile["body"])
            stats.paragraph_count += 1
        return
    if block_type == "internal_heading":
        text = str(block.get("text") or "").strip()
        if text:
            paragraph = _add_body_paragraph(doc, export_profile)
            paragraph.paragraph_format.first_line_indent = None
            run = paragraph.add_run(text)
            run.bold = True
            _format_paragraph_font(paragraph, export_profile["body"])
            stats.paragraph_count += 1
        return
    if block_type == "rich_table":
        _render_rich_table(doc, block, stats, export_profile)
        return
    if block_type == "image_placeholder":
        _render_image_placeholder(doc, block, stats, export_profile)
        return
    if block_type == "image_ref":
        _render_image_group(doc, [block], resolver, stats, profile, export_profile)


def _render_rich_table(
    doc: DocumentObject,
    block: dict[str, Any],
    stats: RenderStats,
    export_profile: dict[str, Any],
) -> None:
    title = str(block.get("title") or "").strip()
    if title:
        caption = doc.add_paragraph()
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = caption.add_run(title)
        run.bold = True
        _format_paragraph_font(caption, export_profile["table"], bold=True)
    columns = [column for column in block.get("columns") or [] if isinstance(column, dict)]
    rows = [row for row in block.get("rows") or [] if isinstance(row, dict)]
    if not columns:
        paragraph = _add_body_paragraph(doc, export_profile)
        paragraph.add_run("【表格列信息缺失，需人工复核】")
        return
    table = doc.add_table(rows=1, cols=len(columns))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    body_width_cm = _body_width_cm(export_profile)
    widths = _column_widths(columns, rows, body_width_cm=body_width_cm, table_profile=export_profile["table"])
    _set_table_width(table, widths)
    for cell, column in zip(table.rows[0].cells, columns):
        cell.text = str(column.get("title") or column.get("key") or "")
        _shade_cell(cell, str(export_profile["table"].get("header_background") or "F2F2F2"))
        _format_cell(cell, export_profile, bold=True, center=True)
    keys = [str(column.get("key") or f"col_{index + 1}") for index, column in enumerate(columns)]
    for row_data in rows:
        row = table.add_row()
        _apply_table_row_height(row, export_profile["table"])
        cells = row_data.get("cells") if isinstance(row_data.get("cells"), dict) else {}
        for cell, key in zip(row.cells, keys):
            cell.text = str(cells.get(key) or "")
            _format_cell(cell, export_profile, center=_is_short_column(key, columns))
    _set_table_width(table, widths)
    stats.table_count += 1
    doc.add_paragraph()


def _render_image_group(
    doc: DocumentObject,
    blocks: list[dict[str, Any]],
    resolver: dict[tuple[str, str], bytes],
    stats: RenderStats,
    profile: ImageLayoutProfile,
    export_profile: dict[str, Any],
) -> None:
    started = time.monotonic()
    stats.image_ref_count += len(blocks)
    try:
        layout_items = [
            ImageLayoutItem(
                block=block,
                image_bytes=_resolve_image_bytes(block, resolver),
                max_columns=_image_max_columns(block, profile),
            )
            for block in blocks
        ]
        if len(layout_items) == 1:
            item = layout_items[0]
            block, image_bytes = item.block, item.image_bytes
            if image_bytes:
                try:
                    _add_picture_paragraph(
                        doc,
                        image_bytes,
                        max_width_cm=profile.single_image_max_width_cm,
                        max_height_cm=profile.single_image_max_height_cm,
                        source_part_name=_image_source_part_name(block),
                    )
                    _add_picture_caption(doc, str(block.get("caption") or "施工做法示意"), export_profile)
                    stats.rendered_image_count += 1
                except UnrecognizedImageError:
                    _render_missing_image_note(doc, block, export_profile)
                    stats.missing_image_count += 1
            else:
                _render_missing_image_note(doc, block, export_profile)
                stats.missing_image_count += 1
            return
        if _should_render_text_image_block_table(blocks):
            _render_text_image_block_image_table(doc, layout_items, blocks, stats, profile, export_profile)
            return
        block_table_title = _text_image_block_table_title(blocks)
        if block_table_title:
            _render_text_image_block_table_title(doc, block_table_title, export_profile)
        group_caption = _shared_image_group_caption(blocks)
        repeated_caption = _repeated_image_group_caption(blocks)
        for row_items in _image_layout_rows(layout_items):
            row_column_count = len(row_items)
            _record_image_layout_row(stats, row_column_count)
            row_image_width_cm = _grid_row_image_max_width_cm(row_column_count, profile)
            row_image_height_cm = _grid_row_image_max_height_cm(row_column_count, profile)
            table = doc.add_table(rows=1, cols=row_column_count)
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            table.style = "Table Grid"
            _set_table_width(table, [profile.body_width_cm / row_column_count] * row_column_count)
            row = table.rows[0]
            for cell_index, layout_item in enumerate(row_items):
                cell = row.cells[cell_index]
                _render_image_layout_item_in_cell(
                    cell,
                    layout_item,
                    max_width_cm=row_image_width_cm,
                    max_height_cm=row_image_height_cm,
                    stats=stats,
                    group_caption=group_caption,
                    repeated_caption=repeated_caption,
                    export_profile=export_profile,
                )
            _set_table_width(table, [profile.body_width_cm / row_column_count] * row_column_count)
        if group_caption:
            _add_picture_caption(doc, group_caption, export_profile)
        doc.add_paragraph()
    finally:
        stats.image_processing_duration_seconds += time.monotonic() - started


def _render_text_image_block_image_table(
    doc: DocumentObject,
    layout_items: list[ImageLayoutItem],
    blocks: list[dict[str, Any]],
    stats: RenderStats,
    profile: ImageLayoutProfile,
    export_profile: dict[str, Any],
) -> None:
    image_rows = _image_layout_rows(layout_items)
    if not image_rows:
        return
    max_columns = max(1, max(len(row) for row in image_rows))
    table = doc.add_table(rows=1, cols=max_columns)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    widths = [profile.body_width_cm / max_columns] * max_columns
    _set_table_width(table, widths)

    has_content_row = False
    title = _text_image_block_table_title(blocks)
    if title:
        title_cell = _merge_row_cells(table.rows[0])
        title_cell.text = title
        _shade_cell(title_cell, "EAF2F8")
        _format_cell(title_cell, export_profile, bold=True, center=True)
        has_content_row = True

    group_caption = _shared_image_group_caption(blocks)
    repeated_caption = _repeated_image_group_caption(blocks)
    for row_items in image_rows:
        row = table.add_row() if has_content_row else table.rows[0]
        has_content_row = True
        _render_image_table_row(
            row,
            row_items,
            max_columns=max_columns,
            stats=stats,
            profile=profile,
            group_caption=group_caption,
            repeated_caption=repeated_caption,
            export_profile=export_profile,
        )

    _set_table_width(table, widths)
    if group_caption:
        _add_picture_caption(doc, group_caption, export_profile)
    doc.add_paragraph()


def _render_image_table_row(
    row: Any,
    row_items: list[ImageLayoutItem],
    *,
    max_columns: int,
    stats: RenderStats,
    profile: ImageLayoutProfile,
    group_caption: str,
    repeated_caption: str,
    export_profile: dict[str, Any],
) -> None:
    row_column_count = len(row_items)
    _record_image_layout_row(stats, row_column_count)
    row_image_width_cm = _grid_row_image_max_width_cm(row_column_count, profile)
    row_image_height_cm = _grid_row_image_max_height_cm(row_column_count, profile)
    cells = _image_row_cells(row, row_column_count=row_column_count, max_columns=max_columns)
    for cell, layout_item in zip(cells, row_items):
        _render_image_layout_item_in_cell(
            cell,
            layout_item,
            max_width_cm=row_image_width_cm,
            max_height_cm=row_image_height_cm,
            stats=stats,
            group_caption=group_caption,
            repeated_caption=repeated_caption,
            export_profile=export_profile,
        )
    for cell in row.cells[row_column_count:]:
        if cell not in cells:
            cell.text = ""
            _format_cell(cell, export_profile, center=True)


def _image_row_cells(row: Any, *, row_column_count: int, max_columns: int) -> list[Any]:
    if row_column_count == 1 and max_columns > 1:
        return [_merge_row_cells(row)]
    return list(row.cells[:row_column_count])


def _merge_row_cells(row: Any) -> Any:
    cell = row.cells[0]
    for next_cell in row.cells[1:]:
        cell = cell.merge(next_cell)
    return cell


def _render_image_layout_item_in_cell(
    cell: Any,
    layout_item: ImageLayoutItem,
    *,
    max_width_cm: float,
    max_height_cm: float,
    stats: RenderStats,
    group_caption: str,
    repeated_caption: str,
    export_profile: dict[str, Any],
) -> None:
    block, image_bytes = layout_item.block, layout_item.image_bytes
    _format_cell(cell, export_profile, center=True)
    if image_bytes:
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        try:
            _add_picture_run(
                paragraph,
                image_bytes,
                max_width_cm=max_width_cm,
                max_height_cm=max_height_cm,
                source_part_name=_image_source_part_name(block),
            )
            item_caption = _image_item_caption(
                block,
                group_caption=group_caption,
                repeated_caption=repeated_caption,
            )
            if item_caption:
                caption = cell.add_paragraph(item_caption)
                caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                _format_paragraph_font(caption, _caption_profile(export_profile))
            stats.rendered_image_count += 1
        except UnrecognizedImageError:
            cell.text = _missing_image_text(block)
            _format_cell(cell, export_profile, center=True)
            stats.missing_image_count += 1
    else:
        cell.text = _missing_image_text(block)
        _format_cell(cell, export_profile, center=True)
        stats.missing_image_count += 1


def _render_image_placeholder(
    doc: DocumentObject,
    block: dict[str, Any],
    stats: RenderStats,
    export_profile: dict[str, Any],
) -> None:
    title = str(block.get("caption") or "图片待补充")
    reason = str(block.get("reason") or "需结合本项目资料人工补充。")
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(f"【{title}】{reason}")
    run.bold = True
    run.font.color.rgb = RGBColor(0x80, 0x40, 0x00)
    _format_paragraph_font(paragraph, export_profile["body"])
    stats.placeholder_count += 1


def _render_review_items(doc: DocumentObject, items: list[Any], export_profile: dict[str, Any]) -> None:
    review_items = [item for item in items if isinstance(item, dict)]
    if not review_items:
        return
    _add_numbered_heading(doc, "人工复核清单", level=2)
    for item in review_items:
        paragraph = doc.add_paragraph(style="List Bullet")
        paragraph.add_run(f"[{item.get('severity') or 'medium'}] ")
        paragraph.add_run(str(item.get("message") or ""))
        _format_paragraph_font(paragraph, export_profile["body"])


def _build_image_resolver(
    material_library_json: str | Path | None,
    raw_root: str | Path,
) -> dict[tuple[str, str], bytes]:
    sources = _load_image_sources(material_library_json, raw_root)
    resolver: dict[tuple[str, str], bytes] = {}
    for source in sources.values():
        for path in source.paths:
            if not path.exists() or path.suffix.lower() != ".docx":
                continue
            try:
                with zipfile.ZipFile(path) as archive:
                    for name in archive.namelist():
                        if name.startswith("word/media/"):
                            resolver[(source.source_id, name)] = archive.read(name)
                            resolver.setdefault(("", name), archive.read(name))
            except zipfile.BadZipFile:
                continue
    return resolver


def _load_image_layout_profile(
    render_profile_json: str | Path | None,
    export_profile: dict[str, Any] | None = None,
) -> ImageLayoutProfile:
    profile = ImageLayoutProfile()
    body_width_cm = _body_width_cm(export_profile)
    image_profile = export_profile.get("image", {}) if isinstance(export_profile, dict) else {}
    if not render_profile_json:
        return ImageLayoutProfile(
            body_width_cm=body_width_cm,
            single_image_max_width_cm=min(float(image_profile.get("max_width_cm") or profile.single_image_max_width_cm), body_width_cm),
            single_image_max_height_cm=float(image_profile.get("max_height_cm") or profile.single_image_max_height_cm),
            grid_cell_max_width_cm=min(profile.grid_cell_max_width_cm, body_width_cm / 2),
            grid_cell_padding_cm=profile.grid_cell_padding_cm,
            one_column_max_height_cm=float(image_profile.get("max_height_cm") or profile.one_column_max_height_cm),
            two_column_max_height_cm=profile.two_column_max_height_cm,
            three_column_max_height_cm=profile.three_column_max_height_cm,
            high_detail_keywords=profile.high_detail_keywords,
            medium_detail_keywords=profile.medium_detail_keywords,
            photo_keywords=profile.photo_keywords,
        )
    path = Path(render_profile_json)
    if not path.exists():
        return _load_image_layout_profile(None, export_profile)
    try:
        data = json.loads(path.read_text(encoding="utf-8"))
    except json.JSONDecodeError:
        return _load_image_layout_profile(None, export_profile)
    image_layout = data.get("image_layout") if isinstance(data, dict) else None
    if not isinstance(image_layout, dict):
        return _load_image_layout_profile(None, export_profile)
    return ImageLayoutProfile(
        body_width_cm=body_width_cm,
        single_image_max_width_cm=_float_or_default(
            image_layout.get("single_image_max_width_cm"),
            min(float(image_profile.get("max_width_cm") or profile.single_image_max_width_cm), body_width_cm),
        ),
        single_image_max_height_cm=_float_or_default(
            image_layout.get("single_image_max_height_cm"),
            float(image_profile.get("max_height_cm") or profile.single_image_max_height_cm),
        ),
        grid_cell_max_width_cm=_float_or_default(
            image_layout.get("grid_cell_max_width_cm"),
            min(profile.grid_cell_max_width_cm, body_width_cm / 2),
        ),
        grid_cell_padding_cm=_float_or_default(
            image_layout.get("grid_cell_padding_cm"),
            profile.grid_cell_padding_cm,
        ),
        one_column_max_height_cm=_float_or_default(
            image_layout.get("one_column_max_height_cm"),
            float(image_profile.get("max_height_cm") or profile.one_column_max_height_cm),
        ),
        two_column_max_height_cm=_float_or_default(
            image_layout.get("two_column_max_height_cm"),
            profile.two_column_max_height_cm,
        ),
        three_column_max_height_cm=_float_or_default(
            image_layout.get("three_column_max_height_cm"),
            profile.three_column_max_height_cm,
        ),
        high_detail_keywords=_string_list_or_default(
            image_layout.get("high_detail_keywords"),
            profile.high_detail_keywords,
        ),
        medium_detail_keywords=_string_list_or_default(
            image_layout.get("medium_detail_keywords"),
            profile.medium_detail_keywords,
        ),
        photo_keywords=_string_list_or_default(image_layout.get("photo_keywords"), profile.photo_keywords),
    )


def _float_or_default(value: Any, default: float) -> float:
    try:
        return float(value)
    except (TypeError, ValueError):
        return default


def _string_list_or_default(value: Any, default: list[str]) -> list[str]:
    if not isinstance(value, list):
        return default
    result = [str(item) for item in value if str(item).strip()]
    return result or default


def _load_image_sources(material_library_json: str | Path | None, raw_root: str | Path) -> dict[str, ImageSource]:
    if not material_library_json:
        return {}
    path = Path(material_library_json)
    if not path.exists():
        return {}
    data = json.loads(path.read_text(encoding="utf-8"))
    result: dict[str, ImageSource] = {}
    for source in data.get("sources") or []:
        if not isinstance(source, dict):
            continue
        source_id = str(source.get("source_id") or "")
        if not source_id:
            continue
        paths = [_resolve_source_path(Path(item), Path(raw_root)) for item in source.get("source_paths") or []]
        result[source_id] = ImageSource(source_id=source_id, paths=[item for item in paths if item is not None])
    return result


def _resolve_source_path(path: Path, raw_root: Path) -> Path | None:
    raw_root = raw_root if raw_root.is_absolute() else Path.cwd() / raw_root
    raw_text = str(path).strip()
    normalized_text = raw_text.replace("\\", "/")
    candidates: list[Path] = [path]

    if normalized_text.startswith("local://raw/"):
        candidates.append(raw_root / normalized_text.removeprefix("local://raw/"))
    elif normalized_text.startswith("local://"):
        candidates.append(raw_root.parent / normalized_text.removeprefix("local://"))

    normalized_path = Path(normalized_text)
    candidates.append(normalized_path)
    if normalized_text.startswith("data/raw/"):
        candidates.append(raw_root / normalized_text.removeprefix("data/raw/"))
    if not normalized_path.is_absolute():
        candidates.append(Path.cwd() / normalized_path)
        candidates.append(raw_root / normalized_path.name)

    for candidate in candidates:
        if candidate.exists():
            return candidate

    if raw_root.exists():
        matches = list(raw_root.rglob(normalized_path.name))
        if matches:
            return matches[0]
    return None


def _resolve_image_bytes(block: dict[str, Any], resolver: dict[tuple[str, str], bytes]) -> bytes | None:
    part_name = _image_source_part_name(block)
    source_id = _image_source_id(block)
    if not part_name:
        return None
    if source_id:
        return resolver.get((source_id, part_name))
    return resolver.get(("", part_name))


def _image_source_id(block: dict[str, Any]) -> str:
    return str(block.get("source_bid_id") or block.get("source_id") or "").strip()


def _image_source_part_name(block: dict[str, Any]) -> str:
    return str(block.get("source_part_name") or block.get("part_name") or "")


def _take_image_group(blocks: list[dict[str, Any]], start: int) -> tuple[list[dict[str, Any]], int]:
    group: list[dict[str, Any]] = []
    cursor = start
    anchor_key = _image_block_group_key(blocks[start]) if start < len(blocks) else ""
    while cursor < len(blocks) and blocks[cursor].get("type") == "image_ref":
        current_key = _image_block_group_key(blocks[cursor])
        if group and (anchor_key or current_key) and current_key != anchor_key:
            break
        group.append(blocks[cursor])
        cursor += 1
    return group, cursor


def _image_block_group_key(block: dict[str, Any]) -> str:
    text_image_block_id = str(block.get("text_image_block_id") or "").strip()
    if str(block.get("source_reuse_mode") or "") == "text_image_block" and text_image_block_id:
        return f"text_image_block:{text_image_block_id}"
    image_group_id = str(block.get("image_group_id") or "").strip()
    if image_group_id:
        return f"image_group:{image_group_id}"
    return ""


def _shared_image_group_caption(blocks: list[dict[str, Any]]) -> str:
    if len(blocks) < 2:
        return ""
    captions = _image_group_captions(blocks)
    if len(captions) != len(blocks):
        return ""
    if len(set(captions)) == 1:
        return captions[0] if _usable_shared_image_group_caption(captions[0]) else ""
    group_ids = {str(block.get("image_group_id") or "") for block in blocks if block.get("image_group_id")}
    if len(group_ids) == 1:
        group_caption = _normalize_render_caption(
            blocks[0].get("group_title") or blocks[0].get("group_semantic_text")
        )
        if (
            group_caption
            and _usable_shared_image_group_caption(group_caption)
            and sum(1 for caption in captions if caption == group_caption) >= 2
        ):
            return group_caption
    return ""


def _repeated_image_group_caption(blocks: list[dict[str, Any]]) -> str:
    if len(blocks) < 2:
        return ""
    captions = _image_group_captions(blocks)
    return captions[0] if len(captions) == len(blocks) and len(set(captions)) == 1 else ""


def _image_group_captions(blocks: list[dict[str, Any]]) -> list[str]:
    captions = [_normalize_render_caption(block.get("caption")) for block in blocks]
    return [caption for caption in captions if caption]


def _usable_shared_image_group_caption(text: str) -> bool:
    return bool(text) and len(text) <= 32 and not _looks_like_merged_image_caption_title(text)


def _text_image_block_table_title(blocks: list[dict[str, Any]]) -> str:
    if not _should_render_text_image_block_table(blocks):
        return ""
    first = blocks[0]
    for value in _text_image_block_title_candidates(first):
        value = _normalize_render_caption(value)
        if value:
            return value
    return ""


def _text_image_block_title_candidates(block: dict[str, Any]) -> list[Any]:
    candidates: list[Any] = []
    for key in ["text_image_block_title", "group_title", "group_semantic_text"]:
        value = block.get(key)
        if _usable_text_image_block_title(value):
            candidates.append(value)
    render_policy = block.get("render_policy") if isinstance(block.get("render_policy"), dict) else {}
    for key in ["title", "table_title"]:
        value = render_policy.get(key)
        if _usable_text_image_block_title(value):
            candidates.append(value)
    return candidates


def _usable_text_image_block_title(value: Any) -> bool:
    text = _normalize_render_caption(value)
    if not text:
        return False
    if len(text) > 24:
        return False
    if _looks_like_merged_image_caption_title(text):
        return False
    weak_titles = {
        "序号",
        "项目",
        "内容",
        "措施",
        "设计说明",
        "序号；设计说明",
        "序号;设计说明",
        "图片",
        "图示",
        "示意图",
    }
    return text not in weak_titles


def _looks_like_merged_image_caption_title(text: str) -> bool:
    process_markers = [
        "示意",
        "做法",
        "节点",
        "接长",
        "大样",
        "流程",
        "平面",
        "立面",
        "剖面",
        "控制",
        "支设",
        "搭设",
    ]
    hit_count = sum(1 for marker in process_markers if marker in text)
    if hit_count >= 3 and len(text) > 16:
        return True
    if any(separator in text for separator in ["；", ";", "、"]) and len(text) > 18:
        return True
    return False


def _should_render_text_image_block_table(blocks: list[dict[str, Any]]) -> bool:
    if len(blocks) < 2:
        return False
    block_ids = {
        str(block.get("text_image_block_id") or "").strip()
        for block in blocks
        if str(block.get("text_image_block_id") or "").strip()
    }
    if len(block_ids) != 1:
        return False
    return all(str(block.get("source_reuse_mode") or "") == "text_image_block" for block in blocks)


def _render_text_image_block_table_title(
    doc: DocumentObject,
    title: str,
    export_profile: dict[str, Any],
) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    _set_table_width(table, [_body_width_cm(export_profile)])
    cell = table.rows[0].cells[0]
    cell.text = title
    _shade_cell(cell, "EAF2F8")
    _format_cell(cell, export_profile, bold=True, center=True)
    _set_table_width(table, [_body_width_cm(export_profile)])


def _image_item_caption(block: dict[str, Any], *, group_caption: str = "", repeated_caption: str = "") -> str:
    caption = _normalize_render_caption(block.get("caption"))
    if group_caption and caption == group_caption:
        return ""
    if repeated_caption and caption == repeated_caption and not _usable_shared_image_group_caption(caption):
        return ""
    return caption or "施工做法示意"


def _normalize_render_caption(value: Any) -> str:
    return re.sub(r"\s+", "", str(value or "")).strip()


def _image_layout_rows(items: list[ImageLayoutItem]) -> list[list[ImageLayoutItem]]:
    rows: list[list[ImageLayoutItem]] = []
    cursor = 0
    while cursor < len(items):
        target_columns = _image_grid_column_count(items[cursor : cursor + 3])
        if target_columns == 3 and sum(1 for item in items[cursor : cursor + 3] if item.max_columns >= 3) < 3:
            target_columns = 2
        if target_columns == 2 and sum(1 for item in items[cursor : cursor + 2] if item.max_columns >= 2) < 2:
            target_columns = 1
        row = items[cursor : cursor + target_columns]
        cursor += len(row)
        rows.append(row)
    return rows


def _record_image_layout_row(stats: RenderStats, column_count: int) -> None:
    if column_count <= 1:
        stats.image_layout_one_column_row_count += 1
    elif column_count == 2:
        stats.image_layout_two_column_row_count += 1
    else:
        stats.image_layout_three_column_row_count += 1


def _grid_row_image_max_width_cm(column_count: int, profile: ImageLayoutProfile | None = None) -> float:
    profile = profile or ImageLayoutProfile()
    if column_count <= 1:
        return profile.single_image_max_width_cm
    return min(profile.grid_cell_max_width_cm, profile.body_width_cm / column_count - profile.grid_cell_padding_cm)


def _grid_row_image_max_height_cm(column_count: int, profile: ImageLayoutProfile | None = None) -> float:
    profile = profile or ImageLayoutProfile()
    if column_count <= 1:
        return profile.one_column_max_height_cm
    if column_count == 2:
        return profile.two_column_max_height_cm
    return profile.three_column_max_height_cm


def _image_grid_column_count(items_or_count: list[ImageLayoutItem] | int) -> int:
    if isinstance(items_or_count, int):
        image_count = items_or_count
        if image_count >= 6:
            return 3
        return 2
    items = items_or_count
    if not items:
        return 2
    if items[0].max_columns <= 1:
        return 1
    if len(items) >= 3 and all(item.max_columns >= 3 for item in items[:3]):
        return 3
    if len(items) >= 2 and all(item.max_columns >= 2 for item in items[:2]):
        return 2
    return 1


def _image_max_columns(block: dict[str, Any], profile: ImageLayoutProfile | None = None) -> int:
    profile = profile or ImageLayoutProfile()
    text = _image_layout_text(block)
    if any(keyword in text for keyword in profile.high_detail_keywords):
        return 1
    if any(keyword in text for keyword in profile.medium_detail_keywords):
        return 2
    if any(keyword in text for keyword in profile.photo_keywords):
        return 3
    return 2


def _image_layout_text(block: dict[str, Any]) -> str:
    parts: list[str] = []
    for key in ["caption", "semantic_text", "group_title", "group_semantic_text", "source_part_name"]:
        value = block.get(key)
        if value:
            parts.append(str(value))
    candidates = block.get("caption_candidates")
    if isinstance(candidates, list):
        parts.extend(str(item) for item in candidates if item)
    return " ".join(parts)


def _add_picture_paragraph(
    doc: DocumentObject,
    image_bytes: bytes,
    *,
    max_width_cm: float,
    max_height_cm: float,
    source_part_name: str = "",
) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_picture_run(
        paragraph,
        image_bytes,
        max_width_cm=max_width_cm,
        max_height_cm=max_height_cm,
        source_part_name=source_part_name,
    )


def _add_picture_run(
    paragraph: Any,
    image_bytes: bytes,
    *,
    max_width_cm: float,
    max_height_cm: float,
    source_part_name: str = "",
) -> None:
    width_cm, height_cm = _fit_image_size_cm(
        image_bytes,
        max_width_cm=max_width_cm,
        max_height_cm=max_height_cm,
        source_part_name=source_part_name,
    )
    try:
        paragraph.add_run().add_picture(BytesIO(image_bytes), width=Cm(width_cm), height=Cm(height_cm))
    except UnrecognizedImageError:
        if _is_supported_vector_image(source_part_name):
            _add_native_vector_picture_run(
                paragraph,
                image_bytes,
                source_part_name=source_part_name,
                width_cm=width_cm,
                height_cm=height_cm,
            )
            return
        raise


def _fit_image_size_cm(
    image_bytes: bytes,
    *,
    max_width_cm: float,
    max_height_cm: float,
    source_part_name: str = "",
) -> tuple[float, float]:
    if _is_supported_vector_image(source_part_name):
        vector_size = _vector_image_size_cm(image_bytes, source_part_name)
        if vector_size is not None:
            return _fit_size_cm(
                vector_size[0],
                vector_size[1],
                max_width_cm=max_width_cm,
                max_height_cm=max_height_cm,
            )
    try:
        with Image.open(BytesIO(image_bytes)) as image:
            width_px, height_px = image.size
    except (UnidentifiedImageError, OSError):
        return _fallback_vector_or_image_size_cm(max_width_cm=max_width_cm, max_height_cm=max_height_cm)
    return _fit_size_cm(width_px, height_px, max_width_cm=max_width_cm, max_height_cm=max_height_cm)


def _fit_size_cm(
    width: float,
    height: float,
    *,
    max_width_cm: float,
    max_height_cm: float,
) -> tuple[float, float]:
    if width <= 0 or height <= 0:
        return max_width_cm, max_height_cm
    ratio = min(max_width_cm / width, max_height_cm / height)
    return max(width * ratio, 0.1), max(height * ratio, 0.1)


def _fallback_vector_or_image_size_cm(*, max_width_cm: float, max_height_cm: float) -> tuple[float, float]:
    return max_width_cm, min(max_height_cm, max_width_cm * 0.65)


def _is_supported_vector_image(source_part_name: str) -> bool:
    return Path(source_part_name).suffix.lower() in VECTOR_IMAGE_CONTENT_TYPES


def _vector_image_size_cm(image_bytes: bytes, source_part_name: str) -> tuple[float, float] | None:
    suffix = Path(source_part_name).suffix.lower()
    if suffix == ".emf":
        return _emf_size_cm(image_bytes)
    if suffix == ".wmf":
        return _wmf_size_cm(image_bytes)
    return None


def _emf_size_cm(image_bytes: bytes) -> tuple[float, float] | None:
    if len(image_bytes) < 40:
        return None
    try:
        record_type, _header_size = struct.unpack_from("<II", image_bytes, 0)
        if record_type != 1:
            return None
        left, top, right, bottom = struct.unpack_from("<llll", image_bytes, 24)
    except struct.error:
        return None
    width_cm = abs(right - left) * 0.001
    height_cm = abs(bottom - top) * 0.001
    if width_cm <= 0 or height_cm <= 0:
        return None
    return width_cm, height_cm


def _wmf_size_cm(image_bytes: bytes) -> tuple[float, float] | None:
    if len(image_bytes) < 22 or int.from_bytes(image_bytes[:4], "little") != 0x9AC6CDD7:
        return None
    try:
        left, top, right, bottom, inch = struct.unpack_from("<hhhhH", image_bytes, 6)
    except struct.error:
        return None
    if inch <= 0:
        return None
    width_cm = abs(right - left) / inch * 2.54
    height_cm = abs(bottom - top) / inch * 2.54
    if width_cm <= 0 or height_cm <= 0:
        return None
    return width_cm, height_cm


def _add_native_vector_picture_run(
    paragraph: Any,
    image_bytes: bytes,
    *,
    source_part_name: str,
    width_cm: float,
    height_cm: float,
) -> None:
    suffix = Path(source_part_name).suffix.lower()
    content_type = VECTOR_IMAGE_CONTENT_TYPES[suffix]
    story_part = paragraph.part
    package = story_part.package
    partname = package.next_partname(f"/word/media/image%d{suffix}")
    image_part = Part(partname, content_type, image_bytes, package)
    relationship_id = story_part.relate_to(image_part, RT.IMAGE)
    filename = Path(source_part_name).name or f"image{suffix}"
    inline = CT_Inline.new_pic_inline(
        story_part.next_id,
        relationship_id,
        filename,
        Cm(width_cm),
        Cm(height_cm),
    )
    paragraph.add_run()._r.add_drawing(inline)


def _add_picture_caption(doc: DocumentObject, caption: str, export_profile: dict[str, Any]) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = _paragraph_alignment(str(export_profile["image"].get("caption_alignment") or "center"))
    paragraph.paragraph_format.space_before = Pt(float(export_profile["image"].get("caption_space_before_pt") or 0))
    paragraph.paragraph_format.space_after = Pt(float(export_profile["image"].get("caption_space_after_pt") or 0))
    paragraph.add_run(caption)
    _format_paragraph_font(paragraph, _caption_profile(export_profile))


def _render_missing_image_note(doc: DocumentObject, block: dict[str, Any], export_profile: dict[str, Any]) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(_missing_image_text(block))
    run.font.color.rgb = RGBColor(0x80, 0x40, 0x00)
    _format_paragraph_font(paragraph, export_profile["body"])


def _missing_image_text(block: dict[str, Any]) -> str:
    caption = str(block.get("caption") or "图片")
    part_name = str(block.get("source_part_name") or block.get("part_name") or "-")
    return f"【图片源未定位】{caption}（{part_name}）"


def _column_widths(
    columns: list[dict[str, Any]],
    rows: list[dict[str, Any]] | None = None,
    *,
    body_width_cm: float = BODY_WIDTH_CM,
    table_profile: dict[str, Any] | None = None,
) -> list[float]:
    """按列语义和内容长度分配表格列宽。

    编标表格通常有“短标识列 + 长措施列 + 中等指标列”的形态。只看表头长度会把
    “岗位/部门”等短值列撑得过宽，因此这里同时参考列名语义和单元格内容长度。
    """

    if not columns:
        return []
    rows = rows or []
    strategy = str((table_profile or {}).get("column_width_strategy") or "auto_by_content_type")
    if strategy == "balanced":
        return [body_width_cm / len(columns)] * len(columns)
    image_column_width_cm = _float_or_default((table_profile or {}).get("image_column_width_cm"), 5.2)
    specs = [_column_width_spec(column, rows, image_column_width_cm=image_column_width_cm) for column in columns]
    return _fit_column_widths(specs, body_width_cm=body_width_cm)


def _column_width_spec(column: dict[str, Any], rows: list[dict[str, Any]], *, image_column_width_cm: float = 5.2) -> dict[str, float]:
    title = str(column.get("title") or column.get("key") or "")
    key = str(column.get("key") or "")
    sample_texts = _column_sample_texts(key, rows)
    max_text_len = max([_display_text_len(title), *[_display_text_len(text) for text in sample_texts]] or [0])
    avg_text_len = (
        sum(_display_text_len(text) for text in sample_texts) / len(sample_texts)
        if sample_texts
        else _display_text_len(title)
    )
    role = _column_role(title)
    if role == "tiny":
        return {"base": 0.85, "min": 0.75, "floor": 0.65, "max": 1.15, "flex": 0.15}
    if role == "short":
        base = 1.45 if max_text_len <= 8 else 1.9
        return {"base": base, "min": 1.15, "floor": 0.95, "max": 2.6, "flex": 0.35}
    if role == "medium":
        base = 2.7 if max_text_len <= 12 else 3.3
        return {"base": base, "min": 2.0, "floor": 1.55, "max": 4.6, "flex": 0.95}
    if role == "metric":
        base = 3.2 if max_text_len <= 18 else 4.0
        return {"base": base, "min": 2.35, "floor": 1.75, "max": 5.4, "flex": 1.2}
    if role == "image":
        base = min(max(image_column_width_cm, 2.0), 10.0)
        return {"base": base, "min": min(base, 4.0), "floor": 2.6, "max": max(base, 6.5), "flex": 0.8}
    base = 4.8
    if avg_text_len >= 30 or max_text_len >= 45:
        base = 6.2
    elif avg_text_len >= 18 or max_text_len >= 28:
        base = 5.6
    return {"base": base, "min": 3.0, "floor": 2.2, "max": 9.2, "flex": 2.4}


def _column_role(title: str) -> str:
    normalized = title.replace(" ", "").replace("\n", "")
    if normalized in {"序号", "编号", "序", "号"}:
        return "tiny"
    if any(word in normalized for word in ["规格型号", "型号规格", "型号/要求", "规格/要求"]):
        return "medium"
    if any(word in normalized for word in ["图片", "照片", "图示", "示意"]):
        return "image"
    if any(
        word in normalized
        for word in [
            "主要职责",
            "职责描述",
            "工作内容",
            "具体职责",
            "具体措施",
            "控制措施",
            "保障措施",
            "处置措施",
            "应急处置",
            "防控措施",
            "管理措施",
            "技术要求",
            "注意事项",
            "处理意见",
            "维护措施",
            "主要议题",
            "分析与对策",
            "应对及管理措施",
            "所需物资",
            "预防与处置",
        ]
    ):
        return "long"
    if any(
        word in normalized
        for word in [
            "单位",
            "数量",
            "人数",
            "月份",
            "日期",
            "时间",
            "时限",
            "频次",
            "周期",
            "时机",
            "比例",
            "结果",
            "状态",
            "是否",
            "等级",
            "岗位",
            "职务",
            "部门",
            "责任人",
            "主持人",
            "审批层级",
            "到岗时间",
            "进场时间",
            "召开频率",
            "报告时限",
        ]
    ):
        return "short"
    if any(word in normalized for word in ["指标", "目标", "标准", "成果", "要求", "验收"]):
        return "metric"
    if any(
        word in normalized
        for word in [
            "项目",
            "类别",
            "类型",
            "名称",
            "科目",
            "班组",
            "人员",
            "分项",
            "部位",
            "工序",
            "控制点",
            "风险",
            "问题",
            "物资",
            "设备",
            "地点",
            "阶段",
        ]
    ):
        return "medium"
    return "long"


def _column_sample_texts(key: str, rows: list[dict[str, Any]]) -> list[str]:
    result = []
    for row in rows:
        cells = row.get("cells") if isinstance(row.get("cells"), dict) else {}
        value = cells.get(key)
        if value is not None:
            result.append(str(value))
    return result


def _display_text_len(text: str) -> int:
    value = str(text or "")
    ascii_count = sum(1 for char in value if ord(char) < 128)
    return len(value) - ascii_count + (ascii_count + 1) // 2


def _fit_column_widths(specs: list[dict[str, float]], *, body_width_cm: float = BODY_WIDTH_CM) -> list[float]:
    """把列宽约束稳定压入 A4 正文宽度，避免极窄列或负宽。"""

    if not specs:
        return []
    floors = [spec["floor"] for spec in specs]
    mins = [max(spec["min"], spec["floor"]) for spec in specs]
    bases = [min(max(spec["base"], spec["min"]), spec["max"]) for spec in specs]

    if sum(mins) > body_width_cm:
        if sum(floors) >= body_width_cm:
            widths = [body_width_cm * width / sum(floors) for width in floors]
        else:
            widths = floors.copy()
            _distribute_width_extra(widths, specs, body_width_cm - sum(widths), mins)
        return _normalize_width_sum(widths, floors, body_width_cm=body_width_cm)

    widths = bases
    if sum(widths) > body_width_cm:
        widths = _shrink_widths(widths, mins, sum(widths) - body_width_cm)
    elif sum(widths) < body_width_cm:
        _distribute_width_extra(widths, specs, body_width_cm - sum(widths), [spec["max"] for spec in specs])
    return _normalize_width_sum(widths, mins, body_width_cm=body_width_cm)


def _distribute_width_extra(
    widths: list[float],
    specs: list[dict[str, float]],
    extra: float,
    limits: list[float],
) -> None:
    while extra > 0.005:
        candidates = [index for index, width in enumerate(widths) if width < limits[index] - 0.005]
        if not candidates:
            break
        flex_total = sum(specs[index]["flex"] for index in candidates) or len(candidates)
        consumed = 0.0
        for index in candidates:
            share = extra * specs[index]["flex"] / flex_total
            new_width = min(widths[index] + share, limits[index])
            consumed += new_width - widths[index]
            widths[index] = new_width
        if consumed <= 0.005:
            break
        extra -= consumed


def _shrink_widths(widths: list[float], limits: list[float], overflow: float) -> list[float]:
    result = widths.copy()
    while overflow > 0.005:
        shrinkable = [max(0.0, width - limit) for width, limit in zip(result, limits)]
        shrink_total = sum(shrinkable)
        if shrink_total <= 0.005:
            break
        consumed = 0.0
        for index, shrink in enumerate(shrinkable):
            if shrink <= 0:
                continue
            share = overflow * shrink / shrink_total
            new_width = max(result[index] - share, limits[index])
            consumed += result[index] - new_width
            result[index] = new_width
        if consumed <= 0.005:
            break
        overflow -= consumed
    return result


def _normalize_width_sum(
    widths: list[float],
    lower_limits: list[float],
    *,
    body_width_cm: float = BODY_WIDTH_CM,
) -> list[float]:
    result = [max(width, 0.1) for width in widths]
    diff = body_width_cm - sum(result)
    if abs(diff) < 0.01:
        return result
    if diff > 0:
        widest = max(range(len(result)), key=lambda index: result[index])
        result[widest] += diff
        return result
    result = _shrink_widths(result, lower_limits, -diff)
    diff = body_width_cm - sum(result)
    if abs(diff) >= 0.01:
        widest = max(range(len(result)), key=lambda index: result[index] - lower_limits[index])
        result[widest] = max(lower_limits[widest], result[widest] + diff)
    return result


def _is_short_column(key: str, columns: list[dict[str, Any]]) -> bool:
    for column in columns:
        if str(column.get("key") or "") != key:
            continue
        title = str(column.get("title") or "")
        return _column_role(title) in {"tiny", "short"}
    return False


def _set_table_width(table: Any, widths_cm: list[float]) -> None:
    widths_dxa = _widths_to_dxa(widths_cm)
    table_width_dxa = sum(widths_dxa)
    for row in table.rows:
        for cell, width_cm, width_dxa in zip(row.cells, widths_cm, widths_dxa):
            cell.width = Cm(width_cm)
            _set_cell_width(cell, width_dxa)
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    for existing in list(tbl_pr.findall(qn("w:tblW"))):
        tbl_pr.remove(existing)
    tbl_w = OxmlElement("w:tblW")
    tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(table_width_dxa))
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    for existing in list(tbl.findall(qn("w:tblGrid"))):
        tbl.remove(existing)
    tbl_grid = OxmlElement("w:tblGrid")
    for width_dxa in widths_dxa:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width_dxa))
        tbl_grid.append(grid_col)
    tbl.insert(1, tbl_grid)


def _widths_to_dxa(widths_cm: list[float]) -> list[int]:
    if not widths_cm:
        return []
    table_width = int(round(sum(widths_cm) / 2.54 * 1440))
    converted = [int(round(width / 2.54 * 1440)) for width in widths_cm[:-1]]
    converted.append(max(1, table_width - sum(converted)))
    return converted


def _set_cell_width(cell: Any, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    for existing in list(tc_pr.findall(qn("w:tcW"))):
        tc_pr.remove(existing)
    tc_w = OxmlElement("w:tcW")
    tc_pr.append(tc_w)
    tc_w.set(qn("w:type"), "dxa")
    tc_w.set(qn("w:w"), str(width_dxa))


def _format_cell(
    cell: Any,
    export_profile: dict[str, Any] | None = None,
    *,
    bold: bool = False,
    center: bool = False,
) -> None:
    table_profile = (export_profile or merge_word_export_profile(None))["table"]
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    for paragraph in cell.paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else _paragraph_alignment(
            str(table_profile.get("default_alignment") or "justify")
        )
        paragraph.paragraph_format.first_line_indent = Cm(0)
        paragraph.paragraph_format.line_spacing = float(table_profile.get("line_spacing") or 1.35)
        paragraph.paragraph_format.space_after = Pt(0)
        _format_paragraph_font(paragraph, table_profile, bold=bold)


def _apply_table_row_height(row: Any, table_profile: dict[str, Any]) -> None:
    min_height_cm = _float_or_default(table_profile.get("min_row_height_cm"), 0)
    if min_height_cm <= 0:
        return
    row.height = Twips(int(round(min_height_cm / 2.54 * 1440)))


def _format_paragraph_font(paragraph: Any, style_profile: dict[str, Any], bold: bool | None = None) -> None:
    for run in paragraph.runs:
        _apply_run_font(run, style_profile, bold=bold)


def _add_body_paragraph(doc: DocumentObject, export_profile: dict[str, Any] | None = None) -> Any:
    body_profile = (export_profile or merge_word_export_profile(None))["body"]
    paragraph = doc.add_paragraph()
    _apply_paragraph_format(paragraph.paragraph_format, body_profile)
    paragraph.alignment = _paragraph_alignment(str(body_profile.get("alignment") or "justify"))
    return paragraph


def _add_numbered_heading(doc: DocumentObject, text: str, *, level: int) -> Any:
    paragraph = doc.add_heading(text, level=level)
    return paragraph


def _record_heading(stats: RenderStats, level: int) -> None:
    stats.heading_count += 1
    if level <= 1:
        stats.heading1_count += 1
    elif level == 2:
        stats.heading2_count += 1
    else:
        stats.heading3_count += 1


def _normalize_output_mode(output_mode: str) -> str:
    value = str(output_mode or REVIEW_DOCX_MODE).strip().lower()
    if value in {FINAL_DOCX_MODE, "正式版", "final_draft"}:
        return FINAL_DOCX_MODE
    return REVIEW_DOCX_MODE


def _shade_cell(cell: Any, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)
