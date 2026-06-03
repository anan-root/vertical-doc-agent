"""诊断章节/整本技术标 JSON 中的图片引用可渲染性。"""

from __future__ import annotations

import argparse
import json
import sys
from io import BytesIO
from pathlib import Path
from typing import Any

from docx import Document
from docx.image.exceptions import UnrecognizedImageError


ROOT = Path(__file__).resolve().parents[2]
SRC = ROOT / "src"
if str(SRC) not in sys.path:
    sys.path.insert(0, str(SRC))

from construction_bidding_agent.chapter_generator.chapter_docx_renderer import (  # noqa: E402
    DEFAULT_RAW_ROOT,
    _add_picture_run,
    _build_image_resolver,
    _resolve_image_bytes,
)


DEFAULT_INPUT = ROOT / "outputs" / "json" / "full_bid_export_result_first10_quality_recheck.json"
DEFAULT_LIBRARY = ROOT / "outputs" / "json" / "excellent_bid_material_library_with_image_assets.json"
DEFAULT_OUTPUT = ROOT / "outputs" / "reports" / "docx_image_ref_diagnostics_first10.md"


def main() -> int:
    parser = argparse.ArgumentParser(description="诊断 Word 渲染前的图片引用是否可定位、可识别。")
    parser.add_argument("--input-json", default=str(DEFAULT_INPUT), help="章节生成结果或整本导出 JSON。")
    parser.add_argument("--material-library", default=str(DEFAULT_LIBRARY), help="优秀标书素材库 JSON。")
    parser.add_argument("--raw-root", default=str(DEFAULT_RAW_ROOT), help="原始文件根目录。")
    parser.add_argument("--output-report", default=str(DEFAULT_OUTPUT), help="Markdown 诊断报告输出路径。")
    args = parser.parse_args()

    data = _load_json(Path(args.input_json))
    resolver = _build_image_resolver(args.material_library, args.raw_root)
    refs = _collect_image_refs(data)
    diagnostics = [_diagnose_ref(item, resolver) for item in refs]
    report = _render_report(diagnostics, resolver_count=len(resolver))
    output = Path(args.output_report)
    output.parent.mkdir(parents=True, exist_ok=True)
    output.write_text(report, encoding="utf-8")
    print(f"Report: {output.resolve()}")
    print(_summary_line(diagnostics, resolver_count=len(resolver)))
    return 1 if any(item["status"] != "ok" for item in diagnostics) else 0


def _load_json(path: Path) -> dict[str, Any]:
    return json.loads(path.read_text(encoding="utf-8"))


def _collect_image_refs(data: dict[str, Any]) -> list[dict[str, Any]]:
    refs: list[dict[str, Any]] = []

    def walk(value: Any, *, chapter_path: list[str], section_heading: str) -> None:
        if isinstance(value, dict):
            next_chapter_path = chapter_path
            next_section_heading = section_heading
            if isinstance(value.get("chapter_path"), list):
                next_chapter_path = [str(item) for item in value["chapter_path"] if str(item).strip()]
            if value.get("heading"):
                next_section_heading = str(value.get("heading"))
            if value.get("type") == "image_ref":
                refs.append(
                    {
                        "chapter_path": next_chapter_path,
                        "section_heading": next_section_heading,
                        "block": value,
                    }
                )
            for child in value.values():
                walk(child, chapter_path=next_chapter_path, section_heading=next_section_heading)
        elif isinstance(value, list):
            for child in value:
                walk(child, chapter_path=chapter_path, section_heading=section_heading)

    walk(data, chapter_path=[], section_heading="")
    return refs


def _diagnose_ref(item: dict[str, Any], resolver: dict[tuple[str, str], bytes]) -> dict[str, Any]:
    block = item["block"]
    image_bytes = _resolve_image_bytes(block, resolver)
    base = {
        "chapter_path": item["chapter_path"],
        "section_heading": item["section_heading"],
        "caption": block.get("caption") or "",
        "source_bid_id": block.get("source_bid_id") or "",
        "source_part_name": block.get("source_part_name") or block.get("part_name") or "",
        "asset_id": block.get("asset_id") or block.get("image_id") or "",
        "semantic_text": block.get("semantic_text") or "",
        "group_id": block.get("group_id") or "",
        "group_title": block.get("group_title") or "",
    }
    if not image_bytes:
        return {**base, "status": "missing", "byte_count": 0}
    try:
        probe_doc = Document()
        _add_picture_run(
            probe_doc.add_paragraph(),
            image_bytes,
            max_width_cm=6.0,
            max_height_cm=4.0,
            source_part_name=str(base["source_part_name"]),
        )
    except UnrecognizedImageError:
        return {**base, "status": "unrecognized", "byte_count": len(image_bytes)}
    return {**base, "status": "ok", "byte_count": len(image_bytes)}


def _render_report(diagnostics: list[dict[str, Any]], *, resolver_count: int) -> str:
    total = len(diagnostics)
    ok_count = sum(1 for item in diagnostics if item["status"] == "ok")
    missing = [item for item in diagnostics if item["status"] == "missing"]
    unrecognized = [item for item in diagnostics if item["status"] == "unrecognized"]
    lines = [
        "# Word 图片引用诊断报告",
        "",
        f"- 素材库媒体索引数：{resolver_count}",
        f"- 图片引用总数：{total}",
        f"- 可正常渲染：{ok_count}",
        f"- 源文件未定位：{len(missing)}",
        f"- 图片格式无法识别：{len(unrecognized)}",
        "",
    ]
    if missing or unrecognized:
        lines.extend(
            [
                "## 异常图片引用",
                "",
                "| 状态 | 章节 | 小节 | 图片说明 | source_bid_id | source_part_name | asset_id |",
                "|---|---|---|---|---|---|---|",
            ]
        )
        for item in missing + unrecognized:
            lines.append(
                "| "
                + " | ".join(
                    [
                        _escape(str(item["status"])),
                        _escape(" > ".join(item["chapter_path"])),
                        _escape(str(item["section_heading"])),
                        _escape(str(item["caption"])),
                        _escape(str(item["source_bid_id"])),
                        _escape(str(item["source_part_name"])),
                        _escape(str(item["asset_id"])),
                    ]
                )
                + " |"
            )
        lines.append("")
    return "\n".join(lines)


def _summary_line(diagnostics: list[dict[str, Any]], *, resolver_count: int) -> str:
    return (
        f"resolver={resolver_count}, refs={len(diagnostics)}, "
        f"ok={sum(1 for item in diagnostics if item['status'] == 'ok')}, "
        f"missing={sum(1 for item in diagnostics if item['status'] == 'missing')}, "
        f"unrecognized={sum(1 for item in diagnostics if item['status'] == 'unrecognized')}"
    )


def _escape(value: str) -> str:
    return value.replace("|", "\\|").replace("\n", " ")


if __name__ == "__main__":
    raise SystemExit(main())
